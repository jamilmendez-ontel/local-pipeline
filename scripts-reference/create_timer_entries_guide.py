#!/usr/bin/env python3
"""Generate the Timer Activity Entries user guide as a Word document.

Usage:
    cd swift_api_pipeline
    venv/Scripts/python ../scripts-reference/create_timer_entries_guide.py

Output: local-pipeline/docs/Timer_Activity_Entries_Guide.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

BRAND_BLUE = RGBColor(0x15, 0x65, 0xC0)
HEADER_BG = "EEF3FA"
WARNING_BG = "FFFBE6"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    tc_pr.append(shading)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_screenshot_placeholder(doc, label):
    """Add a styled screenshot placeholder paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\u27a1 {label}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    # Add spacing
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(12)
    p_fmt.space_after = Pt(12)


def add_heading(doc, text, level=1):
    """Add a heading with brand color for level 1."""
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = BRAND_BLUE
    return h


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Add a numbered list item."""
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def build_guide():
    doc = Document()

    # -- Title page --
    title = doc.add_heading("Timer Activity Entries", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("User Guide")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # -- Tool Description --
    add_heading(doc, "Tool Description")

    doc.add_paragraph(
        "The Timer Activity Entries system automatically emails each technician "
        "a daily recap of their previous day's timer entries from Swift. It gives "
        "techs visibility into their logged time and lets them self-service "
        "corrections and removals without contacting management."
    )

    doc.add_paragraph("Key features:")
    add_bullet(doc, "Daily Task Summary — aggregated view of time by project, site, and task")
    add_bullet(doc, "Entry Details — full breakdown of every individual timer entry")
    add_bullet(doc, "One-click Edit and Remove via Google Forms")
    add_bullet(doc, "Automatic duplicate detection from Swift sync")
    add_bullet(doc, "Daily reminders for unresolved duplicate entries")

    # -- What is this email? --
    add_heading(doc, "What Is This Email?")

    doc.add_paragraph(
        "Every night after the timer pipeline runs, you'll receive an email called "
        "Timer Activity Entries with a summary of your previous day's timer entries "
        "from Swift. This email helps you review your logged time, spot any issues, "
        "and fix them yourself."
    )

    add_screenshot_placeholder(doc, "SCREENSHOT A — Full email view showing the header, greeting, Daily Task Summary, and top of Entry Details")

    # -- Daily Task Summary --
    add_heading(doc, "Reading the Daily Task Summary")

    doc.add_paragraph(
        "At the top of the email is the Daily Task Summary table. This gives you a "
        "quick snapshot of where your day went — your entries grouped by Project, "
        "Site, and Task."
    )

    add_table(doc,
        ["Column", "What It Shows"],
        [
            ["Project", "The project name"],
            ["Site", "The site name"],
            ["Task", "The task name"],
            ["Entries", "How many individual timer entries you have for that task"],
            ["Total", "The total duration across all entries (e.g. \"2h 15m\")"],
            ["Duplicates", "\u26a0 if the system detected possible duplicates, \u2014 if none"],
        ],
        col_widths=[1.5, 4.5],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Rows with a \u26a0 are highlighted in yellow. These are usually caused by "
        "Swift sync creating multiple snapshots of the same timer — don't worry, "
        "it's not something you did wrong."
    )

    add_screenshot_placeholder(doc, "SCREENSHOT B — Close-up of the Daily Task Summary table (at least one row with \u26a0 and one without)")

    # -- Entry Details --
    add_heading(doc, "Reading the Entry Details")

    doc.add_paragraph(
        "Below the summary is the Entry Details table with every individual timer "
        "entry for the day."
    )

    add_table(doc,
        ["Column", "What It Shows"],
        [
            ["Date", "The date of the entry"],
            ["Project", "Project name"],
            ["Site", "Site name"],
            ["Task", "Task name"],
            ["Start", "Start time"],
            ["End", "End time"],
            ["Duration", "How long the entry is"],
            ["Action", "Edit or Remove buttons"],
        ],
        col_widths=[1.5, 4.5],
    )

    doc.add_paragraph()
    doc.add_paragraph("Each row has two buttons:")
    add_bullet(doc, " — fix the duration of this entry", bold_prefix="Edit")
    add_bullet(doc, " — delete this entry", bold_prefix="Remove")

    add_screenshot_placeholder(doc, "SCREENSHOT C — Close-up of Entry Details table showing rows with Edit and Remove buttons")

    # -- How to Edit --
    add_heading(doc, "How to Edit an Entry")

    doc.add_paragraph(
        "Use Edit when the duration on an entry is wrong (e.g., you forgot to stop "
        "the timer and it logged 8 hours instead of 2)."
    )

    add_numbered(doc, "Find the entry in the Entry Details table")
    add_numbered(doc, " button on that row", bold_prefix="Click Edit")
    add_numbered(doc, "A Google Form opens, pre-filled with the entry details")
    add_numbered(doc, " and a brief ", bold_prefix="Fill in the correct duration")
    # Fix: redo step 4 properly
    p = doc.paragraphs[-1]
    p.clear()
    p.style = doc.styles["List Number"]
    run = p.add_run("Fill in the correct duration")
    run.bold = True
    p.add_run(" and a brief ")
    run2 = p.add_run("reason")
    run2.bold = True
    p.add_run(" for the change")

    add_numbered(doc, "", bold_prefix="Click Submit")

    doc.add_paragraph(
        "Your correction is processed automatically — the original entry in Swift "
        "is untouched, but reports will use the corrected duration."
    )

    add_screenshot_placeholder(doc, "SCREENSHOT D — Edit Google Form showing pre-filled entry details, duration field, and reason field")

    # -- How to Remove --
    add_heading(doc, "How to Remove an Entry")

    doc.add_paragraph(
        "Use Remove when an entry shouldn't exist at all — it's a duplicate, "
        "logged to the wrong task, or was created by mistake."
    )

    add_numbered(doc, "Find the entry in the Entry Details table")
    add_numbered(doc, " button on that row", bold_prefix="Click Remove")
    add_numbered(doc, "A Google Form opens, pre-filled with the entry details")
    add_numbered(doc, " to confirm the removal", bold_prefix="Click Submit")

    doc.add_paragraph(
        "That's it — one click to open, one click to confirm."
    )

    add_screenshot_placeholder(doc, "SCREENSHOT E — Remove Google Form showing pre-filled entry details and Submit button")

    # -- Common Scenarios --
    add_heading(doc, "Common Scenarios")

    add_heading(doc, "I logged time to the wrong task", level=2)
    add_numbered(doc, "Start a new timer in Swift under the correct task")
    add_numbered(doc, " the correct entry and fill in the right duration", bold_prefix="Edit")
    add_numbered(doc, " the wrong entry", bold_prefix="Remove")

    add_heading(doc, "My duration is wrong", level=2)
    doc.add_paragraph(
        "Click Edit on that entry, enter the correct duration and reason, "
        "and submit."
    )

    add_heading(doc, "I see a \u26a0 duplicate flag", level=2)
    doc.add_paragraph(
        "This usually means Swift created two copies of the same timer entry "
        "with different end times. Look at the entries in the detail table — "
        "they'll have the same start time but different durations. Remove the "
        "one that's wrong."
    )

    add_heading(doc, "I didn't do anything about the duplicates", level=2)
    doc.add_paragraph(
        "You'll receive a reminder email the next day for any unresolved "
        "duplicates. Reminders continue daily until the duplicates are resolved. "
        "After 7 days with no action, the system automatically keeps the most "
        "recent entry."
    )

    add_screenshot_placeholder(doc, "SCREENSHOT F (optional) — A reminder email for unresolved duplicates")

    # -- FAQ --
    add_heading(doc, "Frequently Asked Questions")

    faq = [
        ("Do I need to install anything?",
         "No. The email arrives automatically. Edit and Remove open Google Forms in your browser."),
        ("Can I fix entries from older emails?",
         "Yes. The Edit and Remove links stay valid — you can go back to a previous day's email and still use the buttons."),
        ("Will this change my data in Swift?",
         "No. Your original Swift timer data is never modified. Corrections and removals are applied to a separate clean copy used for reporting."),
        ("What if I accidentally remove the wrong entry?",
         "Contact Jamil. The original data is preserved and can be restored."),
    ]

    for q, a in faq:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True
        doc.add_paragraph(f"A: {a}")
        doc.add_paragraph()  # spacer

    # -- Screenshot Reference --
    add_heading(doc, "Screenshot Reference (For Guide Author)")

    doc.add_paragraph(
        "Use the following reference when capturing screenshots to insert into "
        "this guide. Replace each placeholder above with the corresponding image."
    )

    add_table(doc,
        ["Screenshot", "What to Capture", "Insert After"],
        [
            ["A", "Full email — header banner through the start of Entry Details", "\"What Is This Email?\""],
            ["B", "Daily Task Summary table — crop tight, include one \u26a0 row and one clean row", "\"Reading the Daily Task Summary\""],
            ["C", "Entry Details table — 4-5 rows showing all columns with Edit/Remove buttons", "\"Reading the Entry Details\""],
            ["D", "Edit Google Form — pre-filled fields, duration input, reason input", "\"How to Edit an Entry\""],
            ["E", "Remove Google Form — pre-filled fields and Submit button", "\"How to Remove an Entry\""],
            ["F", "Reminder email (optional) — if available from a previous run", "\"I didn't do anything about the duplicates\""],
        ],
        col_widths=[1.0, 3.5, 2.0],
    )

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    output_dir = Path(__file__).resolve().parent.parent / "docs"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "Timer_Activity_Entries_Guide.docx"

    doc = build_guide()
    doc.save(str(output_path))
    print(f"Guide saved to: {output_path}")
